#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
read_inputs.py —— 读投后评价项目的输入资料，产出结构化 data.json

它做三件事，把一个杂乱的项目资料文件夹整理成给 build_report.py 用的数据：
    1) Excel 财务数据（.xlsx 用 openpyxl / .xls 用 xlrd）→ 每张表读成表头 + 行数据
    2) 文字型 PDF / txt / docx 文字材料 → 抽成纯文本片段
    3) 扫描件 PDF（图片型）→ 调 ocr_pdf.py 走 tesseract 中文 OCR

产出 data.json 结构（示意）：
    {
      "report_type": null,          # 由人/AI 判断填 "equity"(股权) 或 "engineering"(工程)
      "project_name": null,         # 待 AI 从材料里提取后回填
      "excels": { "文件名.xlsx": { "Sheet1": {"header": [...], "rows": [[...], ...]} } },
      "texts":  { "某材料.pdf": "抽出的正文..." },
      "raw_files": [ "扫到的所有文件路径" ]
    }

注意：本脚本只负责"把资料读出来放进 json 的原始字段"，不做业务判断。
      项目名、金额归位、report_type 这些"理解层"的活由 AI 读着 json 去填对应字段，
      或人工补。脚本保证"资料都进来了、能被程序读到"。

依赖（本机已装）：openpyxl 或 xlrd、pdftotext（poppler）、tesseract（OCR，见 ocr_pdf.py）
    python-docx 可选（读 .docx 文字材料时用）

用法：
    python3 read_inputs.py /path/to/项目资料文件夹                    # 递归扫描整个文件夹
    python3 read_inputs.py /path/to/项目资料文件夹 -o data.json        # 指定输出
    python3 read_inputs.py a.xlsx b.pdf c.txt                        # 也可直接传若干文件
"""

import os
import sys
import json
import glob

# 复用同目录的 ocr_pdf.py 处理扫描件
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
try:
    from ocr_pdf import pdf_to_text
except Exception:
    pdf_to_text = None  # 没有则退化为只用 pdftotext


# ---------- Excel ----------
def read_excel(path):
    """读 .xlsx / .xls，返回 {sheet名: {"header":[...], "rows":[[...],...]}}。"""
    result = {}
    ext = os.path.splitext(path)[1].lower()
    if ext == ".xlsx":
        try:
            from openpyxl import load_workbook
        except ImportError:
            print("[warn] 未装 openpyxl，跳过 xlsx：pip3 install openpyxl", file=sys.stderr)
            return result
        wb = load_workbook(path, data_only=True, read_only=True)
        for ws in wb.worksheets:
            rows = [[("" if c is None else c) for c in r] for r in ws.iter_rows(values_only=True)]
            rows = [r for r in rows if any(str(c).strip() for c in r)]  # 去空行
            if not rows:
                continue
            result[ws.title] = {"header": [str(c) for c in rows[0]],
                                "rows": [[str(c) for c in r] for r in rows[1:]]}
    elif ext == ".xls":
        try:
            import xlrd
        except ImportError:
            print("[warn] 未装 xlrd，跳过 xls：pip3 install xlrd", file=sys.stderr)
            return result
        wb = xlrd.open_workbook(path)
        for ws in wb.sheets():
            rows = [[ws.cell_value(r, c) for c in range(ws.ncols)] for r in range(ws.nrows)]
            rows = [r for r in rows if any(str(c).strip() for c in r)]
            if not rows:
                continue
            result[ws.name] = {"header": [str(c) for c in rows[0]],
                               "rows": [[str(c) for c in r] for r in rows[1:]]}
    return result


# ---------- 文字材料 ----------
def read_text_file(path):
    """读 txt / md，直接返回内容。"""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def read_docx_text(path):
    """读 .docx 正文文字（需要 python-docx）。"""
    try:
        from docx import Document
    except ImportError:
        print("[warn] 未装 python-docx，跳过 docx：pip3 install python-docx", file=sys.stderr)
        return ""
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 表格文字也拼进来
    for t in doc.tables:
        for row in t.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def read_pdf_text(path):
    """PDF：优先复用 ocr_pdf 的两级策略（文字层直抽→扫描件 OCR）。"""
    if pdf_to_text is not None:
        try:
            return pdf_to_text(path)
        except Exception as e:
            print(f"[warn] PDF 读取失败 {path}：{e}", file=sys.stderr)
            return ""
    # 退化：只有 pdftotext
    import subprocess
    try:
        out = subprocess.run(["pdftotext", "-layout", path, "-"],
                             capture_output=True, timeout=120)
        return out.stdout.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"[warn] pdftotext 失败 {path}：{e}", file=sys.stderr)
        return ""


# ---------- 主流程 ----------
def collect_files(inputs):
    """把传入的文件夹/文件展开成文件列表。"""
    files = []
    for item in inputs:
        if os.path.isdir(item):
            files += glob.glob(os.path.join(item, "**", "*"), recursive=True)
        elif os.path.isfile(item):
            files.append(item)
    return [f for f in files if os.path.isfile(f)]


def build_data(inputs):
    data = {
        "report_type": None,   # "equity" / "engineering"，待 AI 判断后回填
        "project_name": None,  # 待 AI 从材料提取回填
        "excels": {},
        "texts": {},
        "raw_files": [],
    }
    for path in collect_files(inputs):
        name = os.path.basename(path)
        ext = os.path.splitext(path)[1].lower()
        data["raw_files"].append(path)
        try:
            if ext in (".xlsx", ".xls"):
                sheets = read_excel(path)
                if sheets:
                    data["excels"][name] = sheets
            elif ext == ".pdf":
                txt = read_pdf_text(path)
                if txt.strip():
                    data["texts"][name] = txt
            elif ext == ".docx":
                txt = read_docx_text(path)
                if txt.strip():
                    data["texts"][name] = txt
            elif ext in (".txt", ".md"):
                data["texts"][name] = read_text_file(path)
            else:
                print(f"[skip] 未处理类型：{name}", file=sys.stderr)
        except Exception as e:
            print(f"[warn] 读取失败 {name}：{e}", file=sys.stderr)
    return data


def main():
    args = [a for a in sys.argv[1:] if not a.startswith("-")]
    out = "data.json"
    if "-o" in sys.argv:
        out = sys.argv[sys.argv.index("-o") + 1]
        args = [a for a in args if a != out]

    if not args:
        print(__doc__)
        sys.exit(1)

    data = build_data(args)
    with open(out, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"[done] 已写出 {out}")
    print(f"       Excel 表：{len(data['excels'])} 个文件")
    print(f"       文字材料：{len(data['texts'])} 个文件")
    print(f"       原始文件：{len(data['raw_files'])} 个")
    print("       下一步：AI 读 data.json，判断 report_type、提取 project_name 和各字段。")


if __name__ == "__main__":
    main()
